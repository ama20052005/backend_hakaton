from html import escape
from pathlib import Path
from typing import List

from app.models import GeneratedReportFile, ReportFormat, ReportPayload


class ReportExportService:
    """Экспортирует аналитическую справку в DOCX и PDF."""

    def export(
        self,
        payload: ReportPayload,
        output_dir: Path,
        report_format: ReportFormat,
        api_prefix: str,
    ) -> List[GeneratedReportFile]:
        output_dir.mkdir(parents=True, exist_ok=True)

        requested_formats = [report_format]
        if report_format == ReportFormat.BOTH:
            requested_formats = [ReportFormat.DOCX, ReportFormat.PDF]

        files: List[GeneratedReportFile] = []
        for item_format in requested_formats:
            filename = f"report_{payload.report_id}.{item_format.value}"
            file_path = output_dir / filename

            if item_format == ReportFormat.DOCX:
                self._export_docx(payload, file_path)
                content_type = (
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            elif item_format == ReportFormat.PDF:
                self._export_pdf(payload, file_path)
                content_type = "application/pdf"
            else:
                continue

            files.append(
                GeneratedReportFile(
                    format=item_format,
                    filename=filename,
                    download_url=(
                        f"{api_prefix}/reports/{payload.report_id}/download"
                        f"?format={item_format.value}"
                    ),
                    content_type=content_type,
                    size_bytes=file_path.stat().st_size,
                )
            )

        return files

    def _export_docx(self, payload: ReportPayload, file_path: Path) -> None:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError as exc:
            raise RuntimeError(
                "DOCX export requires python-docx. Install project dependencies first."
            ) from exc

        document = Document()
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Arial"
        normal_style.font.size = Pt(10)

        title = document.add_heading(payload.title, level=0)
        title.runs[0].font.name = "Arial"
        document.add_paragraph(payload.subtitle)
        document.add_paragraph(
            f"Дата формирования: {payload.generated_at.strftime('%d.%m.%Y %H:%M')}"
        )

        params_table = document.add_table(rows=0, cols=2)
        params_table.style = "Table Grid"
        for label, value in payload.parameters.items():
            row_cells = params_table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

        document.add_paragraph("")
        document.add_heading("Ключевые показатели", level=1)
        metrics_table = document.add_table(rows=1, cols=2)
        metrics_table.style = "Table Grid"
        metrics_header = metrics_table.rows[0].cells
        metrics_header[0].text = "Показатель"
        metrics_header[1].text = "Значение"
        for metric in payload.summary_metrics:
            row_cells = metrics_table.add_row().cells
            row_cells[0].text = metric.label
            row_cells[1].text = metric.value

        for section in payload.sections:
            if section.heading == "Приложение":
                document.add_page_break()
            document.add_heading(section.heading, level=1)
            for paragraph in section.paragraphs:
                document.add_paragraph(paragraph)
            for table in section.tables:
                document.add_paragraph(table.title)
                doc_table = document.add_table(rows=1, cols=len(table.columns))
                doc_table.style = "Table Grid"
                for idx, column in enumerate(table.columns):
                    doc_table.rows[0].cells[idx].text = column
                for row in table.rows:
                    row_cells = doc_table.add_row().cells
                    for idx, value in enumerate(row):
                        row_cells[idx].text = value

        document.save(file_path)

    def _export_pdf(self, payload: ReportPayload, file_path: Path) -> None:
        try:
            from reportlab.lib import colors
            from reportlab.lib.enums import TA_CENTER
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except ImportError as exc:
            raise RuntimeError(
                "PDF export requires reportlab. Install project dependencies first."
            ) from exc

        font_name, font_path = self._resolve_font()
        if font_path:
            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ReportTitle",
            parent=styles["Title"],
            fontName=font_name,
            alignment=TA_CENTER,
            fontSize=18,
            leading=24,
        )
        heading_style = ParagraphStyle(
            "ReportHeading",
            parent=styles["Heading1"],
            fontName=font_name,
            fontSize=14,
            leading=18,
            spaceBefore=12,
            spaceAfter=8,
        )
        body_style = ParagraphStyle(
            "ReportBody",
            parent=styles["BodyText"],
            fontName=font_name,
            fontSize=10,
            leading=14,
            spaceAfter=6,
        )

        story = [
            Paragraph(escape(payload.title), title_style),
            Spacer(1, 4 * mm),
            Paragraph(escape(payload.subtitle), body_style),
            Paragraph(
                escape(
                    f"Дата формирования: {payload.generated_at.strftime('%d.%m.%Y %H:%M')}"
                ),
                body_style,
            ),
            Spacer(1, 4 * mm),
            Paragraph("Параметры отчета", heading_style),
            self._build_pdf_table(
                [["Параметр", "Значение"]]
                + [[label, value] for label, value in payload.parameters.items()],
                font_name,
                colors,
                Table,
                TableStyle,
            ),
            Spacer(1, 4 * mm),
            Paragraph("Ключевые показатели", heading_style),
            self._build_pdf_table(
                [["Показатель", "Значение"]]
                + [[metric.label, metric.value] for metric in payload.summary_metrics],
                font_name,
                colors,
                Table,
                TableStyle,
            ),
        ]

        for section in payload.sections:
            story.append(Spacer(1, 4 * mm))
            story.append(Paragraph(escape(section.heading), heading_style))
            for paragraph in section.paragraphs:
                story.append(Paragraph(escape(paragraph), body_style))
            for table in section.tables:
                story.append(Paragraph(escape(table.title), body_style))
                story.append(
                    self._build_pdf_table(
                        [table.columns] + table.rows,
                        font_name,
                        colors,
                        Table,
                        TableStyle,
                    )
                )

        doc = SimpleDocTemplate(
            str(file_path),
            pagesize=A4,
            leftMargin=15 * mm,
            rightMargin=15 * mm,
            topMargin=15 * mm,
            bottomMargin=15 * mm,
        )
        doc.build(story)

    def _build_pdf_table(self, rows, font_name, colors, table_cls, table_style_cls):
        table = table_cls(rows, repeatRows=1)
        table.setStyle(
            table_style_cls(
                [
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("LEADING", (0, 0), (-1, -1), 11),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9e7f5")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#8aa5c2")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f8fb")]),
                ]
            )
        )
        return table

    def _resolve_font(self):
        candidates = [
            ("ArialUnicode", Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf")),
            ("Arial", Path("/System/Library/Fonts/Supplemental/Arial.ttf")),
            ("TimesNewRoman", Path("/System/Library/Fonts/Supplemental/Times New Roman.ttf")),
            ("DejaVuSans", Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf")),
            ("DejaVuSans", Path("/usr/share/fonts/dejavu/DejaVuSans.ttf")),
            ("LiberationSans", Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf")),
        ]
        for font_name, font_path in candidates:
            if font_path.exists():
                return font_name, font_path
        return "Helvetica", None


report_export_service = ReportExportService()
